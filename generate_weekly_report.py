#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
量子情报周报 v2.0 - Word文档生成
生成日期: 2026-02-24
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import json

# 设置中文字体
def set_chinese_font(run, font_name='SimSun', font_size=10.5, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold

# 创建文档
doc = Document()

# 设置文档默认字体
doc.styles['Normal'].font.name = 'SimSun'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# ==================== 封面 ====================
title = doc.add_heading('', level=0)
title_run = title.add_run('量子情报周报')
title_run.font.size = Pt(26)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
set_chinese_font(title_run, 'Microsoft YaHei', 26, True)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('Quantum Intelligence Weekly Report')
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
subtitle_run.font.italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_para = doc.add_paragraph()
date_run = date_para.add_run('报告周期: 2026年2月17日 - 2026年2月24日')
set_chinese_font(date_run, 'Microsoft YaHei', 12)
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

version_para = doc.add_paragraph()
version_run = version_para.add_run('版本: v2.0 | 生成时间: 2026-02-24')
set_chinese_font(version_run, 'SimSun', 10)
version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # 空行

# ==================== 数据定义 ====================

# 政策动态
policy_data = [
    {
        "title": "U.S. Commerce Department Revises License Review Policy for Exports",
        "date": "2026-01-20",
        "source": "quantum.gov/BIS",
        "summary": "美国商务部工业安全局发布最终规则，修订对华和澳门出口某些先进计算商品的许可审查政策，加强量子技术出口管制。",
        "url": "https://www.globalpolicywatch.com/2026/01/u-s-commerce-department-revises-license-review-policy/",
        "intel_score": 9.2
    },
    {
        "title": "EU Updates Dual-Use Export Control List 2025",
        "date": "2025-09-08",
        "source": "European Commission",
        "summary": "欧盟委员会通过授权法规更新欧盟两用物项出口管制清单，将量子系统、半导体和高性能电子产品纳入管制范围。",
        "url": "https://policy.trade.ec.europa.eu/news/2025-update-eu-control-list-dual-use-items",
        "intel_score": 8.8
    },
    {
        "title": "DARPA Quantum Benchmarking Initiative Stage B",
        "date": "2025-11-06",
        "source": "DARPA",
        "summary": "DARPA选定11家公司进入量子基准测试计划第二阶段，目标是在十年内建造容错量子计算机。",
        "url": "https://www.darpa.mil/research/programs/quantum-benchmarking-initiative",
        "intel_score": 8.5
    },
    {
        "title": "DARPA FY2026 Funding Allocations",
        "date": "2026-01-20",
        "source": "DARPA",
        "summary": "DARPA发布2026财年资金分配，预算49亿美元，较2025年增长12%，重点支持量子计算研究。",
        "url": "https://www.quantum-australia.com/news/darpas-2026-funding-opportunities",
        "intel_score": 8.0
    },
    {
        "title": "2025 US National Security Strategy Policy Reset",
        "date": "2026-01-12",
        "source": "U.S. Government",
        "summary": "2025年美国国家安全战略以'美国优先'和现实主义原则为中心，加强对量子技术等新兴技术的出口管制。",
        "url": "https://www.alvarezandmarsal.com/thought-leadership/what-the-2025-us-national-security-strategy-means",
        "intel_score": 7.5
    }
]

# 学术动态
academic_data = [
    {
        "title": "Fault-tolerant preparation of arbitrary logical states in four-legged cat code",
        "date": "2026-02-20",
        "source": "arXiv",
        "summary": "研究人员提出了在四腿猫编码中制备任意逻辑态的完整容错框架，为容错量子计算提供新途径。",
        "url": "https://arxiv.org/html/2602.17438v1",
        "intel_score": 9.0
    },
    {
        "title": "A fault-tolerant neutral-atom architecture for universal quantum computing",
        "date": "2026-01",
        "source": "Nature",
        "summary": "Nature发表中性原子容错架构研究，实现对逻辑量子比特块的并行控制，支持通用量子计算。",
        "url": "https://www.nature.com/articles/s41586-025-09848-5",
        "intel_score": 9.5
    },
    {
        "title": "IBM lays out clear path to fault-tolerant quantum computing",
        "date": "2026-02",
        "source": "IBM Research",
        "summary": "IBM发布实现大规模容错量子计算机的清晰、严格、全面的框架，目标2029年实现。",
        "url": "https://www.ibm.com/quantum/blog/large-scale-ftqc",
        "intel_score": 9.0
    },
    {
        "title": "Fault-tolerant quantum computation with polylogarithmic overhead",
        "date": "2026-02",
        "source": "Nature Physics",
        "summary": "研究提出多对数开销的容错量子计算方法，显著降低容错量子计算的资源需求。",
        "url": "https://www.nature.com/articles/s41567-025-03102-5",
        "intel_score": 8.5
    },
    {
        "title": "Developments in superconducting erasure qubits for fault-tolerant computing",
        "date": "2026-01-06",
        "source": "arXiv",
        "summary": "超导擦除量子比特研究进展，探讨了纠错与阈值定理对容错量子计算的重要性。",
        "url": "https://arxiv.org/html/2601.02183v2",
        "intel_score": 8.0
    }
]

# 企业动态
company_data = [
    {
        "title": "Google Quantum AI Shows 13000× Speedup Over Supercomputer",
        "date": "2025-10-22",
        "source": "Google Quantum AI",
        "summary": "Google 65量子比特处理器在复杂物理模拟中比Frontier超级计算机快13000倍，展示可验证量子优势。",
        "url": "https://thequantuminsider.com/2025/10/22/google-quantum-ai-shows-13000x-speedup",
        "intel_score": 9.5
    },
    {
        "title": "IBM Quantum Roadmap 2026: 360 Qubits, 7500 Gates",
        "date": "2026",
        "source": "IBM Quantum",
        "summary": "IBM量子路线图2026年目标：实现360量子比特、7500门操作，为容错量子计算铺平道路。",
        "url": "https://www.ibm.com/roadmaps/quantum/2026/",
        "intel_score": 9.0
    },
    {
        "title": "IonQ Only Quantum Company in 2025 Deloitte Technology Fast 500",
        "date": "2025-11-19",
        "source": "IonQ",
        "summary": "IonQ入选2025德勤科技高成长500强唯一量子计算公司，计划实现200万量子比特系统。",
        "url": "https://www.ionq.com/news/ionq-only-quantum-company-in-2025-deloitte-technology-fast-500",
        "intel_score": 8.5
    },
    {
        "title": "IonQ Acquires Oxford Ionics Trap-on-Chip Technology",
        "date": "2025-06",
        "source": "IonQ",
        "summary": "IonQ整合Oxford Ionics芯片离子阱技术，将单阱量子比特数量大幅提升，加速容错计算路径。",
        "url": "https://s28.q4cdn.com/828571518/files/doc_presentation/2025/25-06-09-Webinar-Presentation.pdf",
        "intel_score": 8.0
    },
    {
        "title": "Quantum Computing Funding Explosive Growth in 2025",
        "date": "2025-10-31",
        "source": "Industry Report",
        "summary": "2025年前5个月量子计算融资已达2024年全年的70%，显示投资者对量子技术的强烈信心。",
        "url": "https://www.spinquanta.com/news-detail/quantum-computing-funding-explosive-growth-strategic-investment-2025",
        "intel_score": 7.5
    }
]

# 中国动态
china_data = [
    {
        "title": "本源司南量子操作系统开放线上下载",
        "date": "2026-02-24",
        "source": "本源量子",
        "summary": "我国首款自主研发量子计算机操作系统'本源司南'正式开放线上下载，成为全球首个开放下载的量子操作系统。",
        "url": "https://www.iyiou.com/briefing/202602241880791",
        "intel_score": 9.5
    },
    {
        "title": "国盾量子预计2025年扭亏为盈",
        "date": "2026-01-27",
        "source": "国盾量子",
        "summary": "国盾量子预计2025年营收3.1亿元，同比增长22.34%，Q4净利润创近五年单季度新高，量子计算领域收入增长显著。",
        "url": "https://www.cls.cn/detail/2292570",
        "intel_score": 9.0
    },
    {
        "title": "中国量子科技赛道融资超25亿元",
        "date": "2026-02-22",
        "source": "财联社",
        "summary": "2025年初至2026年2月，中国量子科技赛道融资金额达25.16亿元，融资事件25起，赛道加速升温。",
        "url": "https://finance.sina.cn/stock/jdts/2026-02-22/detail-inhnsyqe0743507.d.html",
        "intel_score": 8.5
    },
    {
        "title": "潘建伟院士：祖冲之三号刷新超导体系纪录",
        "date": "2026-01-04",
        "source": "中国科大",
        "summary": "2025年'祖冲之三号'与'祖冲之3.2号'两代产品接连突破，刷新超导体系量子计算纪录。",
        "url": "https://news.ustc.edu.cn/info/1056/93692.htm",
        "intel_score": 9.0
    },
    {
        "title": "量超融合项目在成都启动",
        "date": "2026-01-26",
        "source": "中国量子信息",
        "summary": "国内量超融合项目取得重要进展，2026年1月26日成都启动相关项目，推动量子计算与超算融合。",
        "url": "https://www.c114.com.cn/quantum/5285/a1305557.html",
        "intel_score": 8.0
    }
]

# 合并所有数据
all_items = policy_data + academic_data + company_data + china_data
top_10 = sorted(all_items, key=lambda x: x["intel_score"], reverse=True)[:10]

# ==================== 执行摘要 ====================
doc.add_page_break()
h1 = doc.add_heading('', level=1)
h1_run = h1.add_run('一、执行摘要 (高价值情报 TOP 10)')
set_chinese_font(h1_run, 'Microsoft YaHei', 16, True)

summary_text = doc.add_paragraph()
summary_run = summary_text.add_run('本周量子情报监测覆盖政策、学术、企业、竞争态势四大板块，共筛选出20条高价值情报。以下是评分最高的10条情报：')
set_chinese_font(summary_run, 'SimSun', 10.5)

for i, item in enumerate(top_10, 1):
    p = doc.add_paragraph(style='List Number')
    
    # 标题和评分
    title_run = p.add_run(f"[{item['intel_score']}] {item['title']}")
    set_chinese_font(title_run, 'SimSun', 10.5, True)
    
    # 来源和日期
    info_para = doc.add_paragraph()
    info_run = info_para.add_run(f"    来源: {item['source']} | 日期: {item['date']}")
    set_chinese_font(info_run, 'SimSun', 9)
    info_run.font.color.rgb = RGBColor(100, 100, 100)
    info_para.paragraph_format.left_indent = Inches(0.3)
    
    # 摘要
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run(f"    {item['summary']}")
    set_chinese_font(summary_run, 'SimSun', 10)
    summary_para.paragraph_format.left_indent = Inches(0.3)

# ==================== 政策动态 ====================
doc.add_page_break()
h2 = doc.add_heading('', level=1)
h2_run = h2.add_run('二、政策动态 (美欧对华量子政策)')
set_chinese_font(h2_run, 'Microsoft YaHei', 16, True)

for item in sorted(policy_data, key=lambda x: x["intel_score"], reverse=True):
    # 标题
    p = doc.add_paragraph()
    title_run = p.add_run(f"• {item['title']}")
    set_chinese_font(title_run, 'SimSun', 11, True)
    
    # 元信息
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"  来源: {item['source']} | 日期: {item['date']} | 评分: {item['intel_score']}")
    set_chinese_font(meta_run, 'SimSun', 9)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)
    meta_para.paragraph_format.left_indent = Inches(0.2)
    
    # 摘要
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run(f"  {item['summary']}")
    set_chinese_font(summary_run, 'SimSun', 10.5)
    summary_para.paragraph_format.left_indent = Inches(0.2)
    
    # URL
    url_para = doc.add_paragraph()
    url_run = url_para.add_run(f"  链接: {item['url']}")
    set_chinese_font(url_run, 'SimSun', 9)
    url_run.font.color.rgb = RGBColor(0, 0, 255)
    url_run.font.underline = True
    url_para.paragraph_format.left_indent = Inches(0.2)
    
    doc.add_paragraph()  # 空行

# ==================== 学术动态 ====================
doc.add_page_break()
h3 = doc.add_heading('', level=1)
h3_run = h3.add_run('三、学术动态 (论文/突破)')
set_chinese_font(h3_run, 'Microsoft YaHei', 16, True)

for item in sorted(academic_data, key=lambda x: x["intel_score"], reverse=True):
    p = doc.add_paragraph()
    title_run = p.add_run(f"• {item['title']}")
    set_chinese_font(title_run, 'SimSun', 11, True)
    
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"  来源: {item['source']} | 日期: {item['date']} | 评分: {item['intel_score']}")
    set_chinese_font(meta_run, 'SimSun', 9)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)
    meta_para.paragraph_format.left_indent = Inches(0.2)
    
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run(f"  {item['summary']}")
    set_chinese_font(summary_run, 'SimSun', 10.5)
    summary_para.paragraph_format.left_indent = Inches(0.2)
    
    url_para = doc.add_paragraph()
    url_run = url_para.add_run(f"  链接: {item['url']}")
    set_chinese_font(url_run, 'SimSun', 9)
    url_run.font.color.rgb = RGBColor(0, 0, 255)
    url_run.font.underline = True
    url_para.paragraph_format.left_indent = Inches(0.2)
    
    doc.add_paragraph()

# ==================== 企业动态 ====================
doc.add_page_break()
h4 = doc.add_heading('', level=1)
h4_run = h4.add_run('四、企业动态 (融资/合作/产品)')
set_chinese_font(h4_run, 'Microsoft YaHei', 16, True)

for item in sorted(company_data, key=lambda x: x["intel_score"], reverse=True):
    p = doc.add_paragraph()
    title_run = p.add_run(f"• {item['title']}")
    set_chinese_font(title_run, 'SimSun', 11, True)
    
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"  来源: {item['source']} | 日期: {item['date']} | 评分: {item['intel_score']}")
    set_chinese_font(meta_run, 'SimSun', 9)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)
    meta_para.paragraph_format.left_indent = Inches(0.2)
    
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run(f"  {item['summary']}")
    set_chinese_font(summary_run, 'SimSun', 10.5)
    summary_para.paragraph_format.left_indent = Inches(0.2)
    
    url_para = doc.add_paragraph()
    url_run = url_para.add_run(f"  链接: {item['url']}")
    set_chinese_font(url_run, 'SimSun', 9)
    url_run.font.color.rgb = RGBColor(0, 0, 255)
    url_run.font.underline = True
    url_para.paragraph_format.left_indent = Inches(0.2)
    
    doc.add_paragraph()

# ==================== 技术突破 ====================
doc.add_page_break()
h5 = doc.add_heading('', level=1)
h5_run = h5.add_run('五、技术突破 (量子比特/纠错/实用化)')
set_chinese_font(h5_run, 'Microsoft YaHei', 16, True)

tech_items = [item for item in all_items if any(kw in item['title'].lower() + item['summary'].lower() 
                                               for kw in ['error correction', 'fault-tolerant', 'logical qubit', 
                                                         '纠错', '容错', '量子比特'])]
tech_items = sorted(tech_items, key=lambda x: x["intel_score"], reverse=True)[:8]

for item in tech_items:
    p = doc.add_paragraph()
    title_run = p.add_run(f"• {item['title']}")
    set_chinese_font(title_run, 'SimSun', 11, True)
    
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"  来源: {item['source']} | 日期: {item['date']} | 评分: {item['intel_score']}")
    set_chinese_font(meta_run, 'SimSun', 9)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)
    meta_para.paragraph_format.left_indent = Inches(0.2)
    
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run(f"  {item['summary']}")
    set_chinese_font(summary_run, 'SimSun', 10.5)
    summary_para.paragraph_format.left_indent = Inches(0.2)
    
    doc.add_paragraph()

# ==================== 竞争态势 ====================
doc.add_page_break()
h6 = doc.add_heading('', level=1)
h6_run = h6.add_run('六、竞争态势 (中美欧对比)')
set_chinese_font(h6_run, 'Microsoft YaHei', 16, True)

# 美国
us_para = doc.add_paragraph()
us_run = us_para.add_run('美国')
set_chinese_font(us_run, 'SimSun', 12, True)
us_run.font.color.rgb = RGBColor(0, 51, 102)

us_text = doc.add_paragraph()
us_text_run = us_text.add_run('• 政策层面：BIS修订对华出口管制规则，DARPA预算增长12%至49亿美元，量子基准测试计划进入第二阶段。\n'
                              '• 技术层面：Google实现13000倍超算加速，IBM发布2029年容错量子计算路线图。\n'
                              '• 企业层面：IonQ入选德勤高成长500强，行业融资保持强劲增长。')
set_chinese_font(us_text_run, 'SimSun', 10.5)
us_text.paragraph_format.left_indent = Inches(0.2)

doc.add_paragraph()

# 欧洲
eu_para = doc.add_paragraph()
eu_run = eu_para.add_run('欧洲')
set_chinese_font(eu_run, 'SimSun', 12, True)
eu_run.font.color.rgb = RGBColor(0, 51, 102)

eu_text = doc.add_paragraph()
eu_text_run = eu_text.add_run('• 政策层面：欧盟更新两用物项出口管制清单，将量子系统纳入管制范围，推进经济安全工具箱建设。\n'
                              '• 技术层面：Nature发表中性原子容错架构重要研究。\n'
                              '• 战略层面：Quantum Europe Strategy推进混合量子计算设施标准化。')
set_chinese_font(eu_text_run, 'SimSun', 10.5)
eu_text.paragraph_format.left_indent = Inches(0.2)

doc.add_paragraph()

# 中国
cn_para = doc.add_paragraph()
cn_run = cn_para.add_run('中国')
set_chinese_font(cn_run, 'SimSun', 12, True)
cn_run.font.color.rgb = RGBColor(0, 51, 102)

cn_text = doc.add_paragraph()
cn_text_run = cn_text.add_run('• 政策层面：量子科技写入国家战略，潘建伟院士强调工程化突破重要性。\n'
                              '• 技术层面：祖冲之三号刷新超导体系纪录，本源司南操作系统全球首个开放下载。\n'
                              '• 产业层面：国盾量子预计2025年扭亏为盈，量子赛道融资超25亿元，量超融合项目启动。')
set_chinese_font(cn_text_run, 'SimSun', 10.5)
cn_text.paragraph_format.left_indent = Inches(0.2)

doc.add_paragraph()

# 对比分析
analysis_para = doc.add_paragraph()
analysis_run = analysis_para.add_run('对比分析')
set_chinese_font(analysis_run, 'SimSun', 12, True)
analysis_run.font.color.rgb = RGBColor(0, 51, 102)

analysis_text = doc.add_paragraph()
analysis_text_run = analysis_text.add_run('本周监测显示，中美欧三方在量子计算领域呈现不同发展重点：\n'
                                          '• 美国侧重出口管制与技术领先，Google、IBM在硬件性能上保持优势；\n'
                                          '• 欧洲强化政策协调与标准制定，出口管制与两用物项清单更新频繁；\n'
                                          '• 中国加速产业化进程，操作系统开源、量超融合、企业盈利等里程碑事件频发，投融资活跃度显著提升。')
set_chinese_font(analysis_text_run, 'SimSun', 10.5)
analysis_text.paragraph_format.left_indent = Inches(0.2)

# ==================== 页脚 ====================
doc.add_page_break()
footer_para = doc.add_paragraph()
footer_run = footer_para.add_run('---\n本报告由量子情报监测系统自动生成\n报告周期: 2026-02-17 至 2026-02-24\n生成时间: 2026-02-24')
set_chinese_font(footer_run, 'SimSun', 9)
footer_run.font.color.rgb = RGBColor(128, 128, 128)
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存文档
output_file = 'quantum_weekly_2026-02-17_2026-02-24.docx'
doc.save(output_file)
print(f"Word文档已生成: {output_file}")
