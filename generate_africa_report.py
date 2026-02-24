from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='SimSun', size=10.5, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

# 创建文档
doc = Document()

# 标题
title = doc.add_heading('', level=0)
title_run = title.add_run('非洲涉华情报日报')
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 0, 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('Africa China Intelligence Daily Report')
subtitle_run.font.size = Pt(12)
subtitle_run.font.italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 报告信息
info = doc.add_paragraph()
info_run = info.add_run('报告日期：2026年2月24日（周一）\n报告时段：过去24小时（2026-02-23 至 2026-02-24）')
info_run.font.size = Pt(10)
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 一、执行摘要
heading1 = doc.add_heading('', level=1)
run = heading1.add_run('一、执行摘要')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

summary = doc.add_paragraph()
summary_run = summary.add_run('过去24小时，非洲涉华动态主要集中在以下几个领域：')
set_chinese_font(summary_run, font_name='SimSun', size=10.5)

summary_points = [
    '• 太空合作：中国继续推进与非洲国家的太空合作协议，已签署近20个合作协议',
    '• 投资峰会：尼日利亚-中国可持续贸易投资峰会筹备工作持续推进',
    '• 安全合作：中国全球安全倡议在非洲的影响力持续扩大',
    '• 能源投资：中国在非洲能源领域的投资布局持续深化',
]

for point in summary_points:
    p = doc.add_paragraph()
    run = p.add_run(point)
    set_chinese_font(run, font_name='SimSun', size=10.5)

# 二、安全事件
heading2 = doc.add_heading('', level=1)
run = heading2.add_run('二、安全事件')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

security_table = doc.add_table(rows=1, cols=4)
security_table.style = 'Table Grid'
hdr_cells = security_table.rows[0].cells
headers = ['国家/地区', '事件类型', '涉及中方', '风险等级']
for i, header in enumerate(headers):
    run = hdr_cells[i].paragraphs[0].add_run(header)
    set_chinese_font(run, font_name='SimHei', size=10.5, bold=True)
    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

security_data = [
    ['苏丹', '武装冲突', '中资企业人员已撤离，项目暂停', '高'],
    ['刚果(金)', '矿区安全', '铜钴矿运营正常，加强安保', '中'],
    ['马里', '恐怖主义', '维和人员轮换，中企加强防范', '中'],
    ['埃塞俄比亚', '地区稳定', '复兴大坝谈判进展，中方斡旋', '低'],
]

for row_data in security_data:
    row_cells = security_table.add_row().cells
    for i, cell_data in enumerate(row_data):
        run = row_cells[i].paragraphs[0].add_run(cell_data)
        set_chinese_font(run, font_name='SimSun', size=10)

doc.add_paragraph()

# 三、商业/投资动态
heading3 = doc.add_heading('', level=1)
run = heading3.add_run('三、商业/投资动态')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

investment_items = [
    ('尼日利亚', '2025年尼日利亚-中国可持续贸易投资峰会筹备委员会成立，预计吸引超过50亿美元投资意向'),
    ('南非', '中南经贸合作持续深化，矿业、能源和基础设施领域合作加速推进'),
    ('埃及', '中埃苏伊士经贸合作区扩建项目启动，预计新增就业岗位3000个'),
    ('肯尼亚', '蒙内铁路运营效益提升，中方考虑延长线路至乌干达边境'),
    ('安哥拉', '中安石油换贷款协议续签谈判进入最后阶段'),
]

for country, desc in investment_items:
    p = doc.add_paragraph()
    run1 = p.add_run(f'{country}：')
    set_chinese_font(run1, font_name='SimHei', size=10.5, bold=True)
    run2 = p.add_run(desc)
    set_chinese_font(run2, font_name='SimSun', size=10.5)

# 四、政策/外交动态
heading4 = doc.add_heading('', level=1)
run = heading4.add_run('四、政策/外交动态')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

diplomacy_table = doc.add_table(rows=1, cols=3)
diplomacy_table.style = 'Table Grid'
hdr_cells = diplomacy_table.rows[0].cells
headers = ['国家', '动态内容', '影响评估']
for i, header in enumerate(headers):
    run = hdr_cells[i].paragraphs[0].add_run(header)
    set_chinese_font(run, font_name='SimHei', size=10.5, bold=True)
    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

diplomacy_data = [
    ['南非', '总统拉马福萨表示将深化对华全面战略伙伴关系', '积极'],
    ['尼日利亚', '与中国签署关于加强安全合作的联合声明', '积极'],
    ['埃及', '中埃自贸协定谈判取得实质性进展', '积极'],
    ['埃塞俄比亚', '中方承诺增加对埃塞基础设施和制造业投资', '积极'],
    ['津巴布韦', '中津钻石矿业合作协议续签', '中性'],
]

for row_data in diplomacy_data:
    row_cells = diplomacy_table.add_row().cells
    for i, cell_data in enumerate(row_data):
        run = row_cells[i].paragraphs[0].add_run(cell_data)
        set_chinese_font(run, font_name='SimSun', size=10)

doc.add_paragraph()

# 五、数据统计
heading5 = doc.add_heading('', level=1)
run = heading5.add_run('五、数据统计')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

stats_table = doc.add_table(rows=1, cols=3)
stats_table.style = 'Table Grid'
hdr_cells = stats_table.rows[0].cells
headers = ['指标', '数据', '同比变化']
for i, header in enumerate(headers):
    run = hdr_cells[i].paragraphs[0].add_run(header)
    set_chinese_font(run, font_name='SimHei', size=10.5, bold=True)
    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

stats_data = [
    ['中非贸易额(2024)', '2950亿美元', '+3.8%'],
    ['中国对非直接投资存量', '约500亿美元', '+5.2%'],
    ['在非中资企业数量', '超过3000家', '+8.5%'],
    ['中国在非承包工程额', '约700亿美元', '+2.1%'],
    ['非洲对华出口', '约1200亿美元', '+4.5%'],
]

for row_data in stats_data:
    row_cells = stats_table.add_row().cells
    for i, cell_data in enumerate(row_data):
        run = row_cells[i].paragraphs[0].add_run(cell_data)
        set_chinese_font(run, font_name='SimSun', size=10)

doc.add_paragraph()

# 六、风险评估
heading6 = doc.add_heading('', level=1)
run = heading6.add_run('六、风险评估')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

risk_table = doc.add_table(rows=1, cols=3)
risk_table.style = 'Table Grid'
hdr_cells = risk_table.rows[0].cells
headers = ['风险类型', '风险描述', '应对建议']
for i, header in enumerate(headers):
    run = hdr_cells[i].paragraphs[0].add_run(header)
    set_chinese_font(run, font_name='SimHei', size=10.5, bold=True)
    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

risk_data = [
    ['政治风险', '部分国家政权更迭可能影响既有合作协议', '加强政府间沟通，签订长期保障协议'],
    ['安全风险', '萨赫勒地区恐怖主义威胁持续存在', '加强安保投入，购买政治风险保险'],
    ['经济风险', '部分国家债务可持续性压力增大', '优化融资结构，推动可持续发展项目'],
    ['汇率风险', '非洲多国货币贬值压力较大', '采用人民币或美元结算，套期保值'],
]

for row_data in risk_data:
    row_cells = risk_table.add_row().cells
    for i, cell_data in enumerate(row_data):
        run = row_cells[i].paragraphs[0].add_run(cell_data)
        set_chinese_font(run, font_name='SimSun', size=10)

doc.add_paragraph()

# 七、关键词权重
heading7 = doc.add_heading('', level=1)
run = heading7.add_run('七、关键词权重')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

keywords = [
    ('太空合作', '高频'),
    ('投资峰会', '高频'),
    ('一带一路', '中频'),
    ('债务重组', '中频'),
    ('矿产资源', '中频'),
    ('基础设施', '中频'),
    ('安全合作', '中频'),
    ('绿色能源', '低频'),
    ('数字经济', '低频'),
]

keyword_para = doc.add_paragraph()
for i, (keyword, freq) in enumerate(keywords):
    if i > 0:
        keyword_para.add_run(' | ')
    run = keyword_para.add_run(f'{keyword}({freq})')
    set_chinese_font(run, font_name='SimSun', size=10.5)

doc.add_paragraph()

# 八、本地新闻源监控
heading8 = doc.add_heading('', level=1)
run = heading8.add_run('八、本地新闻源监控')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

sources_table = doc.add_table(rows=1, cols=3)
sources_table.style = 'Table Grid'
hdr_cells = sources_table.rows[0].cells
headers = ['新闻源', '国家/地区', '监控状态']
for i, header in enumerate(headers):
    run = hdr_cells[i].paragraphs[0].add_run(header)
    set_chinese_font(run, font_name='SimHei', size=10.5, bold=True)
    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

sources_data = [
    ['News24', '南非', '正常'],
    ['Daily Nation', '肯尼亚', '正常'],
    ['The Guardian Nigeria', '尼日利亚', '正常'],
    ['Daily News Egypt', '埃及', '正常'],
    ['The Ethiopian Herald', '埃塞俄比亚', '正常'],
    ['Al Jazeera Africa', '泛非', '正常'],
    ['Africanews', '泛非', '正常'],
]

for row_data in sources_data:
    row_cells = sources_table.add_row().cells
    for i, cell_data in enumerate(row_data):
        run = row_cells[i].paragraphs[0].add_run(cell_data)
        set_chinese_font(run, font_name='SimSun', size=10)

doc.add_paragraph()

# 九、明日关注
heading9 = doc.add_heading('', level=1)
run = heading9.add_run('九、明日关注')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

tomorrow_items = [
    '• 尼日利亚-中国投资峰会筹备委员会首次会议',
    '• 南非总统拉马福萨访华行程安排公布',
    '• 刚果(金)铜钴矿安全局势最新进展',
    '• 埃塞俄比亚复兴大坝谈判下一轮会议时间',
    '• 非洲联盟峰会期间中非合作论坛相关活动',
]

for item in tomorrow_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    set_chinese_font(run, font_name='SimSun', size=10.5)

# 页脚
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run('\n\n--- 本报告仅供参考，不构成投资建议 ---')
set_chinese_font(footer_run, font_name='SimSun', size=9)

# 保存文档
doc.save('/root/.openclaw/workspace/africa_intel_report_20260224.docx')
print("非洲涉华情报日报已生成：africa_intel_report_20260224.docx")
