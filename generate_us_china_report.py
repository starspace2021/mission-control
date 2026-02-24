from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = docx.oxml.OxmlElement(tag)
                tcPr.append(element)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), 'auto')

# 创建文档
doc = Document()

# 设置中文字体
def set_chinese_font(run, font_name='SimSun', size=10.5, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

# 标题
title = doc.add_heading('', level=0)
title_run = title.add_run('美国对华政策监控日报')
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 0, 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('US-China Policy Monitoring Daily Report')
subtitle_run.font.size = Pt(12)
subtitle_run.font.italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 报告信息
info = doc.add_paragraph()
info_run = info.add_run('报告日期：2026年2月24日 10:30 (Asia/Shanghai)\n报告时段：过去24小时（2026-02-23 10:30 至 2026-02-24 10:30）')
info_run.font.size = Pt(10)
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # 空行

# 一、实时告警
heading1 = doc.add_heading('', level=1)
run = heading1.add_run('一、实时告警')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

alert_box = doc.add_paragraph()
alert_run = alert_box.add_run('【重要】特朗普政府维持对华10%额外关税政策')
alert_run.font.size = Pt(11)
alert_run.font.bold = True
alert_run.font.color.rgb = RGBColor(255, 0, 0)

content = doc.add_paragraph()
content_run = content.add_run('根据白宫2025年2月1日发布的公告，特朗普总统依据《国际紧急经济权力法》(IEEPA)对所有中国进口商品加征10%的额外关税。该政策以遏制芬太尼流入和打击非法移民为由，目前仍在持续执行中。')
set_chinese_font(content_run, font_name='SimSun', size=10.5)

# 二、主题要点
heading2 = doc.add_heading('', level=1)
run = heading2.add_run('二、主题要点')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

table1 = doc.add_table(rows=1, cols=3)
table1.style = 'Table Grid'
hdr_cells = table1.rows[0].cells
headers = ['主题', '内容摘要', '影响程度']
for i, header in enumerate(headers):
    run = hdr_cells[i].paragraphs[0].add_run(header)
    set_chinese_font(run, font_name='SimHei', size=10.5, bold=True)
    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加数据行
data = [
    ['关税政策', '特朗普政府维持对华10%额外关税，涉及所有中国进口商品', '高'],
    ['出口管制', '中国于2025年2月4日宣布对五种关键金属实施出口管制', '高'],
    ['半导体限制', '美国继续收紧对华先进半导体制造设备和软件工具的出口管制', '中'],
    ['实体清单', '2025年1月2日，中国商务部将28家美国防务公司列入出口管制清单', '中'],
]

for row_data in data:
    row_cells = table1.add_row().cells
    for i, cell_data in enumerate(row_data):
        run = row_cells[i].paragraphs[0].add_run(cell_data)
        set_chinese_font(run, font_name='SimSun', size=10)

doc.add_paragraph()

# 三、市场影响
heading3 = doc.add_heading('', level=1)
run = heading3.add_run('三、市场影响')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

impact_items = [
    ('股市反应', '中美科技股波动加剧，半导体板块承压'),
    ('汇率波动', '人民币汇率在关税政策影响下呈现短期波动'),
    ('大宗商品', '关键金属出口管制引发全球供应链担忧，钨、碲、铟、钼、铋价格上行'),
    ('贸易流向', '部分制造业订单向东南亚、墨西哥转移趋势持续'),
]

for title, desc in impact_items:
    p = doc.add_paragraph()
    run1 = p.add_run(f'{title}：')
    set_chinese_font(run1, font_name='SimHei', size=10.5, bold=True)
    run2 = p.add_run(desc)
    set_chinese_font(run2, font_name='SimSun', size=10.5)

# 四、政策趋势
heading4 = doc.add_heading('', level=1)
run = heading4.add_run('四、政策趋势')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

trends = [
    '1. 关税政策持续性：特朗普政府明确表示关税将"持续有效，直到危机缓解"，短期内难以取消。',
    '2. 技术脱钩深化：美国继续在半导体、人工智能等关键技术领域加强对华出口管制。',
    '3. 中方反制常态化：中国采取"以牙还牙"策略，对关键矿产和特定美国实体实施对等反制。',
    '4. 多边博弈加剧：美国试图联合盟友构建对华技术封锁网络，中国则加强与全球南方国家合作。',
]

for trend in trends:
    p = doc.add_paragraph()
    run = p.add_run(trend)
    set_chinese_font(run, font_name='SimSun', size=10.5)

# 五、今日关注
heading5 = doc.add_heading('', level=1)
run = heading5.add_run('五、今日关注')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

focus_items = [
    '• 美国财政部是否发布新的对华制裁名单',
    '• 中国商务部对关键矿产出口管制的具体实施细则',
    '• 中美双方是否就贸易问题进行非正式接触',
    '• 欧盟对华贸易政策动向及其对中美博弈的影响',
]

for item in focus_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    set_chinese_font(run, font_name='SimSun', size=10.5)

# 六、走势研判
heading6 = doc.add_heading('', level=1)
run = heading6.add_run('六、走势研判')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

judgment = doc.add_paragraph()
judge_run = judgment.add_run('短期（1-3个月）：')
set_chinese_font(judge_run, font_name='SimHei', size=10.5, bold=True)
judge_content = judgment.add_run('关税政策预计将维持现状，双方博弈聚焦于执行层面。技术出口管制清单可能进一步扩大，中方反制措施将保持对等原则。')
set_chinese_font(judge_content, font_name='SimSun', size=10.5)

judgment2 = doc.add_paragraph()
judge_run2 = judgment2.add_run('中期（3-6个月）：')
set_chinese_font(judge_run2, font_name='SimHei', size=10.5, bold=True)
judge_content2 = judgment2.add_run('若芬太尼和边境问题取得进展，关税政策存在调整窗口。但结构性技术竞争将持续，半导体和AI领域脱钩趋势难逆转。')
set_chinese_font(judge_content2, font_name='SimSun', size=10.5)

# 七、数据源监控
heading7 = doc.add_heading('', level=1)
run = heading7.add_run('七、数据源监控')
set_chinese_font(run, font_name='SimHei', size=14, bold=True)

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr_cells2 = table2.rows[0].cells
headers2 = ['数据源', '监控内容', '更新频率']
for i, header in enumerate(headers2):
    run = hdr_cells2[i].paragraphs[0].add_run(header)
    set_chinese_font(run, font_name='SimHei', size=10.5, bold=True)
    hdr_cells2[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

source_data = [
    ['白宫官网', '总统行政令、事实清单', '实时'],
    ['美国财政部OFAC', '制裁名单更新', '每日'],
    ['美国商务部BIS', '出口管制条例修订', '每日'],
    ['中国商务部', '反制措施公告', '实时'],
    ['美国贸易代表办公室', '301调查进展', '每周'],
    ['联邦公报', '法规正式文本', '每日'],
]

for row_data in source_data:
    row_cells = table2.add_row().cells
    for i, cell_data in enumerate(row_data):
        run = row_cells[i].paragraphs[0].add_run(cell_data)
        set_chinese_font(run, font_name='SimSun', size=10)

# 页脚
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run('\n\n--- 本报告仅供参考，不构成投资建议 ---')
set_chinese_font(footer_run, font_name='SimSun', size=9)

# 保存文档
doc.save('/root/.openclaw/workspace/us_china_policy_report_20260224.docx')
print("美国对华政策监控日报已生成：us_china_policy_report_20260224.docx")
